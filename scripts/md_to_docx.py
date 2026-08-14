"""
Convert TRIAL_MANUAL.md (markdown) to a staff-distributable .docx.

Handles the constructs used by the manual: H1-H3, fenced code blocks,
pipe tables, bullet/numbered lists, blockquotes, inline **bold** and
`code` spans, and horizontal rules.

    python scripts/md_to_docx.py docs/TRIAL_MANUAL.md docs/TRIAL_MANUAL.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches

MONO = "Consolas"
BODY = "Calibri"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)


def add_runs(par, text: str) -> None:
    """Render inline **bold** / `code` spans into a paragraph."""
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|`([^`]+)`", text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            r = par.add_run(m.group(1))
            r.bold = True
        else:
            r = par.add_run(m.group(2))
            r.font.name = MONO
            r.font.size = Pt(9.5)
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    cells = [c for c in rows if not all(re.fullmatch(r":?-{2,}:?", x.strip()) for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, val.strip())
            if i == 0:
                for r in p.runs:
                    r.bold = True
    doc.add_paragraph()


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY
    style.font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    table: list[list[str]] = []
    in_code = False
    code_buf: list[str] = []

    def close_table() -> None:
        nonlocal table
        if table:
            flush_table(doc, table)
            table = []

    for line in lines:
        s = line.rstrip()

        if s.startswith("```"):
            close_table()
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r = p.add_run("\n".join(code_buf))
                r.font.name = MONO
                r.font.size = Pt(8.5)
                code_buf = []
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(line)
            continue

        if s.startswith("|") and s.endswith("|"):
            table.append([c.strip() for c in s.strip("|").split("|")])
            continue
        close_table()

        if not s.strip():
            continue
        if re.fullmatch(r"-{3,}", s.strip()):
            continue
        if s.startswith("# "):
            h = doc.add_heading(level=1)
            add_runs(h, s[2:])
            for r in h.runs:
                r.font.color.rgb = ACCENT
        elif s.startswith("## "):
            h = doc.add_heading(level=2)
            add_runs(h, s[3:])
            for r in h.runs:
                r.font.color.rgb = ACCENT
        elif s.startswith("### "):
            h = doc.add_heading(level=3)
            add_runs(h, s[4:])
        elif s.startswith("> "):
            p = doc.add_paragraph()
            add_runs(p, s[2:])
            for r in p.runs:
                r.italic = True
        elif re.match(r"^\s*[-*] ", s):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*] ", "", s))
        elif re.match(r"^\s*\d+\. ", s):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\s*\d+\. ", "", s))
        else:
            p = doc.add_paragraph()
            add_runs(p, s)
    close_table()
    doc.save(str(docx_path))


if __name__ == "__main__":
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "docs/TRIAL_MANUAL.md")
    dst = Path(sys.argv[2] if len(sys.argv) > 2 else src.with_suffix(".docx"))
    convert(src, dst)
    print(f"written: {dst}")
